"""
Generateur de PDF ordonnance - page unique, codes-barres corriges.
"""
import subprocess
import os
import shutil
import tempfile
import base64
from datetime import date
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image
import io
import logging

logger = logging.getLogger(__name__)

TEMPLATE_DIR = Path(__file__).parent.parent / "template"
LOGO_PATH = TEMPLATE_DIR / "logo_clinique.jpg"
SIG_PATH = TEMPLATE_DIR / "signature_clinique.png"


def barcode_png(value, bar_w=3, height=50):
    """Genere un code-barres CODE128B en PNG base64."""
    C128 = [
        "11011001100","11001101100","11001100110","10010011000","10010001100",
        "10001001100","10011001000","10011000100","10001100100","11001001000",
        "11001000100","11000100100","10110011100","10011011100","10011001110",
        "10111001100","10011101100","10011100110","11001110010","11001011100",
        "11001001110","11011100100","11001110100","11101101110","11101001100",
        "11100101100","11100100110","11101100100","11100110100","11100110010",
        "11011011000","11011000110","11000110110","10100011000","10001011000",
        "10001000110","10110001000","10001101000","10001100010","11010001000",
        "11000101000","11000100010","10110111000","10110001110","10001101110",
        "10111011000","10111000110","10001110110","11101110110","11010001110",
        "11000101110","11011101000","11011100010","11011101110","11101011000",
        "11101000110","11100010110","11101101000","11101100010","11100011010",
        "11101111010","11001000010","11110001010","10100110000","10100001100",
        "10010110000","10010000110","10000101100","10000100110","10110010000",
        "10110000100","10011010000","10011000010","10000110100","10000110010",
        "11000010010","11001010000","11110111010","11000010100","10001111010",
        "10100111100","10010111100","10010011110","10111100100","10011110100",
        "10011110010","11110100100","11110010100","11110010010","11011011110",
        "11011110110","11110110110","10101111000","10100011110","10001011110",
        "10111101000","10111100010","11110101000","11110100010","10111011110",
        "10111101110","11101011110","11110101110","11010000100","11010010000",
        "11010011100","11000111010"
    ]
    START_B = 104
    STOP = 106
    chars = [START_B]
    chk = START_B
    for i, c in enumerate(value):
        code = ord(c) - 32
        chars.append(code)
        chk += (i + 1) * code
    chars.append(chk % 103)
    chars.append(STOP)
    bits = ""
    for c in chars:
        if c < len(C128):
            bits += C128[c]
    quiet = 10
    total_w = len(bits) * bar_w + 2 * quiet
    img = Image.new("RGB", (total_w, height), "white")
    px = img.load()
    x = quiet
    for bit in bits:
        color = (0, 0, 0) if bit == "1" else (255, 255, 255)
        for dx in range(bar_w):
            for dy in range(height):
                px[x + dx, dy] = color
        x += bar_w
    buf = io.BytesIO()
    img.save(buf, "PNG")
    return buf.getvalue()


def generate_pdf(
    patient_name,
    patient_dob,
    prescriptions,
    output_dir="/tmp",
    doctor_name="Benjamin BONNOT",
    doctor_specialty="Anesthesiste-Reanimateur",
    doctor_rpps="751031329",
    clinic_name="Clinique Moussins-Nollet",
    clinic_address="67 rue de Romainville, 75019 PARIS",
    clinic_phone="01 40 03 12 12",
    clinic_finess="750301160",
    clinic_rpps_code="10100661908",
):
    today = _fmt_today()
    docx_path = os.path.join(output_dir, "ordonnance.docx")
    _build_docx(
        docx_path=docx_path,
        patient_name=patient_name,
        patient_dob=patient_dob,
        prescriptions=prescriptions,
        today=today,
        doctor_name=doctor_name,
        doctor_specialty=doctor_specialty,
        doctor_rpps=doctor_rpps,
        clinic_address=clinic_address,
        clinic_phone=clinic_phone,
        clinic_finess=clinic_finess,
        clinic_rpps_code=clinic_rpps_code,
    )
    return _to_pdf(docx_path, output_dir, patient_name)


def _fmt_today():
    months = ["janvier","fevrier","mars","avril","mai","juin",
              "juillet","aout","septembre","octobre","novembre","decembre"]
    d = date.today()
    return str(d.day) + " " + months[d.month - 1] + " " + str(d.year)


def _build_docx(docx_path, patient_name, patient_dob, prescriptions, today,
                doctor_name, doctor_specialty, doctor_rpps,
                clinic_address, clinic_phone, clinic_finess, clinic_rpps_code):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(1.8)
    sec.right_margin = Cm(1.8)
    sec.top_margin = Cm(1.2)
    sec.bottom_margin = Cm(1.5)

    # Generer les codes-barres
    bc_finess = barcode_png(clinic_finess, bar_w=2, height=40)
    bc_rpps = barcode_png(clinic_rpps_code, bar_w=2, height=40)

    # En-tete 3 colonnes
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    _set_col_width(tbl.columns[0], Cm(4.5))
    _set_col_width(tbl.columns[1], Cm(9.5))
    _set_col_width(tbl.columns[2], Cm(4.5))

    # Cellule gauche: logo + adresse
    cl = tbl.rows[0].cells[0]
    if LOGO_PATH.exists():
        p = cl.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(LOGO_PATH), width=Cm(2.8))
    addr_parts = clinic_address.split(",")
    addr_p = cl.add_paragraph()
    addr_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    addr_r = addr_p.add_run(addr_parts[0] + "\n" + (addr_parts[1].strip() if len(addr_parts) > 1 else "") + "\nTel : " + clinic_phone)
    addr_r.font.size = Pt(7)
    addr_r.font.name = "Arial"

    # Cellule centre: docteur
    cm = tbl.rows[0].cells[1]
    _clear_cell(cm)
    p1 = cm.add_paragraph("Docteur " + doctor_name)
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.runs[0].bold = True
    p1.runs[0].font.size = Pt(11)
    p1.runs[0].font.name = "Arial"
    p2 = cm.add_paragraph("drbonnot@gmail.com")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.runs[0].font.size = Pt(8)
    p2.runs[0].font.name = "Arial"
    for txt in [doctor_specialty, doctor_rpps]:
        pp = cm.add_paragraph(txt)
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp.runs[0].font.size = Pt(8)
        pp.runs[0].font.name = "Arial"

    # Cellule droite: codes-barres
    cr = tbl.rows[0].cells[2]
    _clear_cell(cr)
    for label, code, bc_data in [
        ("Code FINESS", clinic_finess, bc_finess),
        ("Code RPPS", clinic_rpps_code, bc_rpps),
    ]:
        lp = cr.add_paragraph(label)
        lp.runs[0].font.size = Pt(7)
        lp.runs[0].font.name = "Arial"
        bp = cr.add_paragraph()
        bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = bp.add_run()
        run.add_picture(io.BytesIO(bc_data), width=Cm(3.8))
        cp = cr.add_paragraph(code)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].font.size = Pt(7)
        cp.runs[0].font.name = "Courier New"
        cr.add_paragraph()

    doc.add_paragraph()

    # Nom patient
    pt_p = doc.add_paragraph()
    pt_r = pt_p.add_run(patient_name.upper())
    pt_r.bold = True
    pt_r.font.size = Pt(11)
    pt_r.font.name = "Arial"
    if patient_dob:
        dob_p = doc.add_paragraph()
        dob_r = dob_p.add_run("Ne(e) le : " + patient_dob)
        dob_r.font.size = Pt(10)
        dob_r.font.name = "Arial"

    doc.add_paragraph()

    # Titre ORDONNANCE
    ord_p = doc.add_paragraph()
    ord_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ord_r = ord_p.add_run("ORDONNANCE")
    ord_r.bold = True
    ord_r.underline = True
    ord_r.font.size = Pt(14)
    ord_r.font.name = "Arial"

    doc.add_paragraph()

    # Traitements
    for i, med in enumerate(prescriptions):
        denom = med.get("denomination", "")
        ligne = med.get("ligne", "")
        # Extraire posologie (tout sauf le nom)
        posologie = ligne.replace(denom, "").strip().lstrip("\n").strip()

        med_p = doc.add_paragraph()
        med_p.paragraph_format.space_before = Pt(3)
        med_p.paragraph_format.space_after = Pt(1)
        num_r = med_p.add_run(str(i + 1) + ". ")
        num_r.bold = True
        num_r.font.size = Pt(10)
        num_r.font.name = "Arial"
        name_r = med_p.add_run(denom)
        name_r.bold = True
        name_r.font.size = Pt(10)
        name_r.font.name = "Arial"

        if posologie:
            for line in posologie.split("\n"):
                line = line.strip()
                if line:
                    pos_p = doc.add_paragraph()
                    pos_p.paragraph_format.left_indent = Cm(0.8)
                    pos_p.paragraph_format.space_before = Pt(0)
                    pos_p.paragraph_format.space_after = Pt(1)
                    pos_r = pos_p.add_run(line)
                    pos_r.font.size = Pt(10)
                    pos_r.font.name = "Arial"

        doc.add_paragraph()

    # Pied de page
    for _ in range(3):
        doc.add_paragraph()

    date_p = doc.add_paragraph("Paris, le " + today)
    date_p.runs[0].font.size = Pt(10)
    date_p.runs[0].font.name = "Arial"

    # Signature image
    sig_p = doc.add_paragraph()
    sig_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if SIG_PATH.exists():
        sig_run = sig_p.add_run()
        sig_run.add_picture(str(SIG_PATH), width=Cm(3.5))
    dr_p = doc.add_paragraph()
    dr_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    dr_r = dr_p.add_run("Docteur " + doctor_name + "\n" + doctor_specialty)
    dr_r.font.size = Pt(10)
    dr_r.font.name = "Arial"

    doc.save(docx_path)
    logger.info("Docx genere: " + docx_path)


def _to_pdf(docx_path, output_dir, patient_name):
    safe = patient_name.upper().replace(" ", "_").replace("/", "-")
    today_str = date.today().strftime("%Y%m%d")
    pdf_name = "Ordonnance_" + safe + "_" + today_str + ".pdf"
    pdf_out = os.path.join(output_dir, pdf_name)
    for cmd in ["libreoffice", "soffice", "/usr/bin/libreoffice", "/usr/bin/soffice"]:
        if shutil.which(cmd):
            result = subprocess.run(
                [cmd, "--headless", "--convert-to", "pdf", "--outdir", output_dir, docx_path],
                capture_output=True, text=True, timeout=60,
            )
            if result.returncode == 0:
                generated = os.path.join(output_dir, os.path.basename(docx_path).replace(".docx", ".pdf"))
                if os.path.exists(generated):
                    os.rename(generated, pdf_out)
                return pdf_out
    raise RuntimeError("LibreOffice non trouve - impossible de generer le PDF")


def _set_col_width(col, width):
    for cell in col.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement("w:tcW")
        tcW.set(qn("w:w"), str(int(width.twips)))
        tcW.set(qn("w:type"), "dxa")
        tcPr.append(tcW)


def _clear_cell(cell):
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    cell.paragraphs[0].clear()
